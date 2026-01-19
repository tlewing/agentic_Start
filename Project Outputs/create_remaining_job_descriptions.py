import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

output_dir = 'C:/Users/tewing/Desktop/Claude Projects/Revised Job Descriptions'
os.makedirs(output_dir, exist_ok=True)

def create_revised_job_description(jd_data, filename):
    """Create a revised Job Description with embedded cross-references"""
    doc = Document()

    # Title
    title = doc.add_heading(jd_data['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Policy Reference
    p = doc.add_paragraph()
    p.add_run(f"Policy Manual Reference: {jd_data['policy_ref']}").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Cross-Reference Summary Table
    doc.add_heading('Cross-Reference Summary', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Reference Type'
    hdr_cells[1].text = 'References'

    for ref_type, refs in jd_data['cross_refs'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = ref_type
        row_cells[1].text = refs

    doc.add_paragraph()

    # Position Summary
    doc.add_heading('Position Summary', level=1)
    doc.add_paragraph(jd_data['summary'])

    # Reports To
    doc.add_heading('Reports To', level=1)
    doc.add_paragraph(jd_data['reports_to'])

    # Key Responsibilities
    doc.add_heading('Key Responsibilities', level=1)
    for resp in jd_data['responsibilities']:
        doc.add_paragraph(resp, style='List Bullet')

    # Applicable SOPs Section
    doc.add_heading('Applicable Standard Operating Procedures', level=1)
    doc.add_paragraph('This position is responsible for compliance with the following SOPs:')
    for sop_category, sops in jd_data['applicable_sops'].items():
        doc.add_heading(sop_category, level=2)
        for sop in sops:
            doc.add_paragraph(sop, style='List Bullet')

    # Required Training Section
    doc.add_heading('Required Training', level=1)
    doc.add_paragraph('Personnel in this position must complete the following training:')
    for category, trainings in jd_data['required_training'].items():
        doc.add_heading(category, level=2)
        for training in trainings:
            doc.add_paragraph(training, style='List Bullet')

    # Management Directive Compliance
    doc.add_heading('Management Directive Compliance', level=1)
    doc.add_paragraph('This position operates under the following Management Directives:')
    for md in jd_data['management_directives']:
        doc.add_paragraph(md, style='List Bullet')

    # Qualifications
    doc.add_heading('Qualifications', level=1)
    for qual in jd_data['qualifications']:
        doc.add_paragraph(qual, style='List Bullet')

    # Career Progression
    doc.add_heading('Career Progression', level=1)
    doc.add_paragraph(jd_data['progression'])

    # Save
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f'Created: {filename}')

# ============================================
# REMAINING JOB DESCRIPTIONS
# ============================================

print('Creating Remaining Revised Job Descriptions with Cross-References...')
print('=' * 60)

# LEAD JOURNEYMAN (7.2.3)
lead_journeyman = {
    'title': 'Lead Journeyman Electrician',
    'policy_ref': 'Policy Manual 7.2.3',
    'cross_refs': {
        'Management Directives': 'MD 9.2 (Field Employees); MD 9.3 (Field Leadership)',
        'Primary SOPs': '9.41.470-510 (Safety); 9.41.675-725 (Field Operations)',
        'Required Training': 'Leadership Fundamentals; All Journeyman Training',
        'Target Skill Rules': 'Lead Journeyman Leadership; Lead Journeyman Coordination',
        'Skill Level Sets': 'Leadership & Field Management (Intermediate)'
    },
    'summary': 'The Lead Journeyman assists the Foreman in supervising work crews, provides technical guidance to Journeymen and Apprentices, and ensures quality workmanship on assigned tasks. This position serves as a bridge between field workers and field supervision.',
    'reports_to': 'Foreman or Project Superintendent',
    'responsibilities': [
        'Assist Foreman with crew supervision and task assignment',
        'Provide technical guidance and mentoring to Journeymen and Apprentices',
        'Ensure quality standards are maintained on assigned work areas',
        'Coordinate with other trades as directed by Foreman',
        'Monitor crew productivity and report issues to Foreman',
        'Assist with material staging and tool management',
        'Conduct informal safety observations and correct unsafe behaviors',
        'Step in for Foreman during absences when designated',
        'Complete accurate time records for assigned crew members',
        'Communicate job progress and issues to supervision'
    ],
    'applicable_sops': {
        'Safety SOPs': [
            '9.41.470 Site-Specific Safety Plan Compliance',
            '9.41.475 Safety Orientation Participation',
            '9.41.480 Safety Inspection Support',
            '9.41.485 Safety Meeting Participation',
            '9.41.490-510 Incident Response Support'
        ],
        'Field Operations SOPs': [
            '9.41.675 Field Work Coordination',
            '9.41.680 Quality Workmanship Standards',
            '9.41.685 Trade Coordination Support',
            '9.41.690 Material Handling',
            '9.41.695 Tool Control'
        ],
        'Documentation SOPs': [
            '9.41.215 Daily Report Input Support',
            '9.41.220 Time Recording Assistance'
        ]
    },
    'required_training': {
        'Safety Training': [
            'All Journeyman Safety Training',
            'Safety Leadership Introduction',
            'Hazard Recognition Advanced'
        ],
        'Leadership Training': [
            'Leadership Introduction (SCORM)',
            'Team Building Fundamentals',
            'Mentoring & Coaching Basics',
            'Field Supervision Introduction'
        ],
        'Technical Training': [
            'All Journeyman Technical Training',
            'Quality Control Fundamentals'
        ]
    },
    'management_directives': [
        'MD 9.2 - Field Employee General Requirements',
        'MD 9.3 - Field Leadership Requirements (applicable sections)'
    ],
    'qualifications': [
        'Valid Journeyman Electrician license',
        'Minimum 3 years experience as Journeyman Electrician',
        'Demonstrated leadership ability and positive attitude',
        'Strong technical knowledge across electrical systems',
        'Good communication and interpersonal skills',
        'Willingness to mentor and develop others',
        'Completion of Leadership Introduction training'
    ],
    'progression': 'Lead Journeyman -> Foreman (with additional training and demonstrated leadership capability)'
}

# PROJECT SUPERINTENDENT (7.2.5)
project_superintendent = {
    'title': 'Project Superintendent',
    'policy_ref': 'Policy Manual 7.2.5',
    'cross_refs': {
        'Management Directives': 'MD 9.3 (Field Leadership); MD 9.4 (General Superintendents)',
        'Primary SOPs': 'All 9.2.xxx-9.6.xxx Series (Complete Project Lifecycle)',
        'Required Training': 'All Foreman Training + Advanced Leadership',
        'Target Skill Rules': 'Superintendent Advanced Leadership; Superintendent Project Oversight; Superintendent Team Development',
        'Skill Level Sets': 'Leadership & Field Management (Expert); Project Planning (Expert); Performance Management (Expert)'
    },
    'summary': 'The Project Superintendent provides overall field leadership for large or multiple projects, coordinates multiple Foremen and crews, manages project schedules and resources, and serves as the primary field contact for clients and general contractors. This position is responsible for project success from mobilization through closeout.',
    'reports_to': 'Operations Manager or Branch Manager',
    'responsibilities': [
        'Oversee all field operations for assigned project(s)',
        'Coordinate and supervise multiple Foremen and crews',
        'Manage project schedule and ensure milestone achievement',
        'Serve as primary field contact for clients and general contractors',
        'Conduct project planning and resource allocation',
        'Monitor and manage project budgets and productivity',
        'Ensure compliance with all safety requirements',
        'Resolve field issues and make critical decisions',
        'Coordinate with Project Manager on project status',
        'Develop and mentor Foremen for career advancement',
        'Conduct performance reviews for field supervision',
        'Represent company at project meetings',
        'Manage change order field implementation',
        'Ensure quality standards are maintained'
    ],
    'applicable_sops': {
        'Pre-Construction SOPs (9.2.xxx)': [
            '9.2.050 Plans Review Leadership',
            '9.2.080 Project Turnover Meeting Participation',
            '9.2.100 Site Visit Coordination',
            '9.2.150 Kickoff Meeting Leadership'
        ],
        'Site Setup SOPs (9.3.xxx)': [
            '9.3.001-9.3.017 All Site Setup Oversight'
        ],
        'Project Execution SOPs (9.41.xxx)': [
            '9.41.025 Meeting Management',
            '9.41.050 Project Communication',
            '9.41.100 Schedule Management',
            '9.41.150 Budget Management',
            '9.41.470-510 Safety Program Oversight',
            '9.41.675-725 Field Operations Management'
        ],
        'Closeout SOPs (9.6.xxx)': [
            '9.6.001-9.6.010 Project Closeout Leadership'
        ]
    },
    'required_training': {
        'Advanced Leadership': [
            'Values-Driven Leadership (SCORM)',
            'Extreme Ownership (SCORM)',
            'EQ for Leadership (SCORM)',
            'Career Development Conversations'
        ],
        'Project Management': [
            'Project Planning Advanced',
            'Scheduling Advanced',
            'Productivity Analysis',
            'ProCore Advanced',
            'Cost Management'
        ],
        'Safety Leadership': [
            'Safety Coordinator Certification (all 11 modules)',
            'Safety Program Management'
        ],
        'Performance Management': [
            'Performance Management Fundamentals',
            'Coaching & Feedback Advanced',
            'Team Development'
        ]
    },
    'management_directives': [
        'MD 9.3 - Field Leadership Requirements (full compliance)',
        'MD 9.4 - General Superintendent Requirements',
        'MD 9.5 - Project Management Requirements (field sections)'
    ],
    'qualifications': [
        'Minimum 5 years experience as Foreman',
        'Demonstrated success managing large or multiple projects',
        'Completion of all Foreman training requirements',
        'Strong leadership and team development skills',
        'Excellent communication and client relations abilities',
        'Advanced knowledge of scheduling and cost control',
        'Safety Coordinator certification or equivalent',
        'Proven track record of project success'
    ],
    'progression': 'Project Superintendent -> Operations Manager -> Branch Manager'
}

# OPERATIONS MANAGER (7.3.2)
operations_manager = {
    'title': 'Operations Manager',
    'policy_ref': 'Policy Manual 7.3.2',
    'cross_refs': {
        'Management Directives': 'MD 9.5 (Project Management); MD 9.6 (Branch Management)',
        'Primary SOPs': 'Strategic oversight of all project SOPs',
        'Required Training': 'All PM Training + Executive Leadership',
        'Target Skill Rules': 'Operations Manager Leadership',
        'Skill Level Sets': 'Professional Development (Expert); Leadership (Expert)'
    },
    'summary': 'The Operations Manager oversees multiple Project Managers and Superintendents, manages branch operations and resources, and ensures consistent execution of company standards across all projects. This position is responsible for operational excellence and team development.',
    'reports_to': 'Branch Manager or Vice President of Operations',
    'responsibilities': [
        'Oversee multiple Project Managers and Project Superintendents',
        'Manage branch operational performance and metrics',
        'Allocate resources across projects',
        'Ensure consistent execution of company standards',
        'Develop and mentor project management team',
        'Review and approve project budgets and schedules',
        'Resolve escalated project issues',
        'Represent company to major clients',
        'Participate in business development activities',
        'Implement operational improvements',
        'Conduct performance reviews for direct reports',
        'Ensure compliance with all company policies'
    ],
    'applicable_sops': {
        'Strategic Oversight': [
            'All 9.2.xxx Pre-Construction SOPs (oversight)',
            'All 9.41.xxx Execution SOPs (oversight)',
            'All 9.6.xxx Closeout SOPs (oversight)'
        ],
        'Management SOPs': [
            'Resource allocation procedures',
            'Performance review procedures',
            'Escalation procedures'
        ]
    },
    'required_training': {
        'Executive Leadership': [
            'Strategic Thinking',
            'Presentation Skills',
            'Industry Leadership',
            'Business Development'
        ],
        'All Project Management Training': [
            'All PM training requirements',
            'Advanced financial management',
            'Contract management'
        ]
    },
    'management_directives': [
        'MD 9.5 - Project Management Requirements (full compliance)',
        'MD 9.6 - Branch Management Requirements'
    ],
    'qualifications': [
        'Minimum 5 years experience as Project Manager',
        'Demonstrated success in multi-project management',
        'Strong financial and business acumen',
        'Excellent leadership and team development skills',
        'Strategic thinking capability',
        'Client relationship management experience'
    ],
    'progression': 'Operations Manager -> Branch Manager -> Vice President'
}

# HR MANAGER (7.4.1)
hr_manager = {
    'title': 'Human Resource Manager',
    'policy_ref': 'Policy Manual 7.4.1',
    'cross_refs': {
        'Management Directives': 'All MDs (compliance oversight)',
        'Primary SOPs': 'HR and compliance-related SOPs',
        'Required Training': 'HR compliance; Performance Management',
        'Target Skill Rules': 'HR Manager Compliance',
        'Skill Level Sets': 'Documentation & Compliance (Expert)'
    },
    'summary': 'The Human Resource Manager oversees all HR functions including recruitment, employee relations, benefits administration, compliance, and training coordination. This position ensures company compliance with employment laws and maintains positive employee relations.',
    'reports_to': 'Branch Manager or Vice President of HR',
    'responsibilities': [
        'Manage recruitment and hiring processes',
        'Oversee employee onboarding and orientation',
        'Administer benefits and compensation programs',
        'Handle employee relations issues',
        'Ensure compliance with employment laws and regulations',
        'Coordinate training and development programs',
        'Maintain personnel records and documentation',
        'Support performance management processes',
        'Investigate workplace complaints',
        'Develop and update HR policies',
        'Manage workers compensation claims',
        'Coordinate with safety department on compliance'
    ],
    'applicable_sops': {
        'HR SOPs': [
            'Recruitment and hiring procedures',
            'Onboarding procedures',
            'Performance review administration',
            'Employee separation procedures',
            'Records management procedures'
        ],
        'Compliance SOPs': [
            'Policy compliance monitoring',
            'Investigation procedures',
            'Documentation standards'
        ]
    },
    'required_training': {
        'HR Compliance': [
            'Employment law fundamentals',
            'Documentation best practices',
            'Investigation procedures',
            'Benefits administration'
        ],
        'Management Skills': [
            'Performance Management Fundamentals',
            'Conflict resolution',
            'Communication skills'
        ]
    },
    'management_directives': [
        'All Management Directives (compliance oversight responsibility)',
        'Specific enforcement of MD 9.2 employee requirements'
    ],
    'qualifications': [
        'Bachelor degree in HR, Business, or related field',
        'PHR or SHRM-CP certification preferred',
        'Minimum 5 years HR experience',
        'Strong knowledge of employment law',
        'Excellent communication and interpersonal skills',
        'Experience with HRIS systems',
        'Construction industry experience preferred'
    ],
    'progression': 'HR Manager -> Director of HR -> VP of Human Resources'
}

# ADMINISTRATIVE SUPPORT (7.4.5)
admin_support = {
    'title': 'Administrative Support',
    'policy_ref': 'Policy Manual 7.4.5',
    'cross_refs': {
        'Management Directives': 'MD 9.5 (support functions)',
        'Primary SOPs': 'Document management; Administrative procedures',
        'Required Training': 'Software skills; Documentation practices',
        'Target Skill Rules': 'Admin Support Foundation',
        'Skill Level Sets': 'Construction Software (Foundational)'
    },
    'summary': 'Administrative Support personnel provide clerical, documentation, and administrative services to support project and office operations. This position is essential for maintaining organized records and efficient office functions.',
    'reports_to': 'Office Manager or Project Manager',
    'responsibilities': [
        'Manage document filing and organization',
        'Process incoming and outgoing correspondence',
        'Maintain office supplies and equipment',
        'Support project documentation needs',
        'Answer phones and direct calls appropriately',
        'Schedule meetings and appointments',
        'Prepare reports and presentations',
        'Process expense reports and invoices',
        'Maintain database entries and records',
        'Coordinate with vendors and suppliers',
        'Support project closeout documentation',
        'Assist with onboarding documentation'
    ],
    'applicable_sops': {
        'Documentation SOPs': [
            'Document management procedures',
            'Filing standards and organization',
            'Records retention requirements',
            'Correspondence standards'
        ],
        'Project Support SOPs': [
            '9.6.xxx Closeout documentation support',
            'Submittal tracking support',
            'RFI logging support'
        ]
    },
    'required_training': {
        'Software Skills': [
            'Microsoft Excel Skills',
            'Microsoft Outlook',
            'Microsoft Teams',
            'ViewPoint basics (if applicable)',
            'ProCore basics (if applicable)'
        ],
        'Professional Development': [
            'Documentation Best Practices',
            'Time Management',
            'Professional Communication'
        ]
    },
    'management_directives': [
        'MD 9.5 - Project Management (support function requirements)'
    ],
    'qualifications': [
        'High school diploma or equivalent',
        'Proficiency in Microsoft Office Suite',
        'Strong organizational skills',
        'Attention to detail',
        'Good communication skills',
        'Ability to prioritize tasks',
        'Previous administrative experience preferred'
    ],
    'progression': 'Administrative Support -> Senior Admin -> Office Manager or Specialized Admin Role'
}

# Create all remaining job descriptions
remaining_jds = [
    (lead_journeyman, 'Policy Manual 7.2.3 Lead Journeyman (Revised).docx'),
    (project_superintendent, 'Policy Manual 7.2.5 Project Superintendent (Revised).docx'),
    (operations_manager, 'Policy Manual 7.3.2 Operations Manager (Revised).docx'),
    (hr_manager, 'Policy Manual 7.4.1 Human Resource Manager (Revised).docx'),
    (admin_support, 'Policy Manual 7.4.5 Administrative Support (Revised).docx'),
]

for jd_data, filename in remaining_jds:
    create_revised_job_description(jd_data, filename)

print()
print('=' * 60)
print('Remaining Revised Job Descriptions created with:')
print('  - Cross-Reference Summary Table')
print('  - Applicable SOPs section')
print('  - Required Training section')
print('  - Management Directive Compliance section')
