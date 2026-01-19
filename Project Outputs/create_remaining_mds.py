import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

output_dir = 'C:/Users/tewing/Desktop/Claude Projects/Revised Management Directives'
os.makedirs(output_dir, exist_ok=True)

def create_revised_md(md_data, filename):
    """Create a revised Management Directive with embedded cross-references"""
    doc = Document()

    # Title
    title = doc.add_heading(md_data['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Policy Reference
    p = doc.add_paragraph()
    p.add_run(f"Policy Manual Reference: {md_data['policy_ref']}").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Cross-Reference Summary Table
    doc.add_heading('Cross-Reference Summary', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Reference Type'
    hdr_cells[1].text = 'References'

    for ref_type, refs in md_data['cross_refs'].items():
        row_cells = table.add_row().cells
        row_cells[0].text = ref_type
        row_cells[1].text = refs

    doc.add_paragraph()

    # Purpose
    doc.add_heading('Purpose', level=1)
    doc.add_paragraph(md_data['purpose'])

    # Scope
    doc.add_heading('Scope', level=1)
    doc.add_paragraph(md_data['scope'])

    # Key Directives
    doc.add_heading('Key Directives', level=1)
    for directive in md_data['directives']:
        doc.add_paragraph(directive, style='List Bullet')

    # Implementing SOPs Section
    doc.add_heading('Implementing Standard Operating Procedures', level=1)
    doc.add_paragraph('The following SOPs implement the requirements of this Management Directive:')
    for category, sops in md_data['implementing_sops'].items():
        doc.add_heading(category, level=2)
        for sop in sops:
            doc.add_paragraph(sop, style='List Bullet')

    # Training Requirements Section
    doc.add_heading('Training Requirements', level=1)
    doc.add_paragraph('Personnel subject to this directive must complete the following training:')
    for category, trainings in md_data['training_requirements'].items():
        doc.add_heading(category, level=2)
        for training in trainings:
            doc.add_paragraph(training, style='List Bullet')

    # Job Description Alignment
    doc.add_heading('Job Description Alignment', level=1)
    doc.add_paragraph('This Management Directive applies to the following positions:')
    for jd in md_data['job_descriptions']:
        doc.add_paragraph(jd, style='List Bullet')

    # Compliance Requirements
    doc.add_heading('Compliance Requirements', level=1)
    for req in md_data['compliance']:
        doc.add_paragraph(req, style='List Bullet')

    # Save
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f'Created: {filename}')

# ============================================
# REMAINING MANAGEMENT DIRECTIVES
# ============================================

print('Creating Remaining Revised Management Directives with Cross-References...')
print('=' * 60)

# MD 9.4 - GENERAL SUPERINTENDENTS
md_9_4 = {
    'title': 'Management Directives for General Superintendents',
    'policy_ref': 'Policy Manual 9.4',
    'cross_refs': {
        'Related MDs': 'MD 9.3 (Field Leadership); MD 9.5 (Project Management)',
        'Implementing SOPs': '9.2.xxx-9.6.xxx (Complete Project Lifecycle)',
        'Target Skill Rules': 'Superintendent Advanced Leadership; Superintendent Project Oversight; Superintendent Team Development',
        'Job Descriptions': '7.2.5 (Project Superintendent)',
        'Skill Level Sets': 'Leadership & Field Management (Expert); Project Planning (Expert); Performance Management (Expert)'
    },
    'purpose': 'This Management Directive establishes the requirements, responsibilities, and expectations for General/Project Superintendents. Superintendents serve as the primary field leaders responsible for overall project execution and success.',
    'scope': 'This directive applies to all personnel serving in Project Superintendent or General Superintendent roles, including those supervising multiple Foremen or managing large-scale projects.',
    'directives': [
        'Superintendents shall maintain complete oversight of all field operations for assigned projects',
        'Superintendents shall ensure compliance with all safety programs and requirements',
        'Superintendents shall coordinate with Project Managers on project status, budgets, and schedules',
        'Superintendents shall develop and mentor Foremen to advance their careers',
        'Superintendents shall serve as primary field contact for clients and general contractors',
        'Superintendents shall conduct regular performance reviews for field supervision',
        'Superintendents shall ensure quality standards are maintained across all work',
        'Superintendents shall manage resources effectively across crews and projects',
        'Superintendents shall resolve field issues promptly and escalate when appropriate',
        'Superintendents shall complete all required training and maintain certifications',
        'Superintendents shall lead by example in safety, quality, and professional conduct'
    ],
    'implementing_sops': {
        'Pre-Construction (9.2.xxx)': [
            '9.2.050 Plans Review - Lead field review of project documents',
            '9.2.080 Project Turnover Meeting - Primary field representative',
            '9.2.100 Site Visit - Conduct and document site assessments',
            '9.2.150 Kickoff Meeting - Lead field portion of kickoff'
        ],
        'Site Setup (9.3.xxx)': [
            '9.3.001-9.3.017 All Site Setup - Oversee complete site mobilization'
        ],
        'Project Execution (9.41.xxx)': [
            '9.41.025 Meetings - Lead field coordination meetings',
            '9.41.050 Communication - Maintain stakeholder communication',
            '9.41.100 Schedule - Manage and update project schedules',
            '9.41.150 Budget - Monitor and manage project budgets',
            '9.41.470-510 Safety - Ensure safety program compliance',
            '9.41.600 Quality - Maintain quality standards',
            '9.41.675-725 Field Operations - Direct all field activities'
        ],
        'Closeout (9.6.xxx)': [
            '9.6.001-9.6.010 - Lead project closeout activities'
        ]
    },
    'training_requirements': {
        'Advanced Leadership (Required)': [
            'Values-Driven Leadership (SCORM)',
            'Extreme Ownership (SCORM)',
            'EQ for Leadership (SCORM)',
            'Career Development Conversations'
        ],
        'Project Management (Required)': [
            'Project Planning Advanced',
            'Scheduling Advanced',
            'Productivity Analysis',
            'ProCore Advanced',
            'Cost Management'
        ],
        'Safety Leadership (Required)': [
            'Safety Coordinator Certification (all 11 modules)',
            'Safety Program Management'
        ],
        'Performance Management (Required)': [
            'Performance Management Fundamentals',
            'Coaching & Feedback Advanced',
            'Team Development'
        ],
        'Prerequisites': [
            'All Foreman training requirements completed'
        ]
    },
    'job_descriptions': [
        '7.2.5 Project Superintendent - Primary role governed by this directive',
        '7.2.4 Foreman - Reports to Superintendent; Superintendent responsible for development',
        '7.3.1 Project Manager - Coordination partner for project success'
    ],
    'compliance': [
        'Complete all required training within 6 months of promotion to Superintendent',
        'Maintain Safety Coordinator certification through recertification',
        'Conduct quarterly performance reviews for all direct reports',
        'Attend monthly Superintendent meetings',
        'Submit weekly project status reports',
        'Complete annual leadership development activities'
    ]
}

# MD 9.6 - BRANCH MANAGEMENT
md_9_6 = {
    'title': 'Management Directives for Branch Management',
    'policy_ref': 'Policy Manual 9.6',
    'cross_refs': {
        'Related MDs': 'MD 9.4 (General Superintendents); MD 9.5 (Project Management)',
        'Implementing SOPs': 'Branch operational procedures; Strategic oversight of all project SOPs',
        'Target Skill Rules': 'Operations Manager Leadership; Branch Manager Leadership',
        'Job Descriptions': '7.3.2 (Operations Manager); Branch Manager',
        'Skill Level Sets': 'Professional Development (Expert); Leadership (Expert)'
    },
    'purpose': 'This Management Directive establishes the requirements, responsibilities, and expectations for Branch Management personnel including Operations Managers and Branch Managers. Branch management is responsible for operational excellence, team development, and business performance.',
    'scope': 'This directive applies to all personnel serving in Operations Manager, Branch Manager, or equivalent management roles with oversight of multiple projects and project teams.',
    'directives': [
        'Branch management shall ensure consistent execution of company standards across all projects',
        'Branch management shall develop and mentor project management and field leadership teams',
        'Branch management shall manage branch operational performance and metrics',
        'Branch management shall allocate resources effectively across projects',
        'Branch management shall represent the company to major clients and stakeholders',
        'Branch management shall participate in business development activities',
        'Branch management shall implement operational improvements and best practices',
        'Branch management shall ensure compliance with all company policies and procedures',
        'Branch management shall conduct performance reviews and career development planning',
        'Branch management shall manage branch financial performance',
        'Branch management shall coordinate with other branches and corporate functions'
    ],
    'implementing_sops': {
        'Strategic Oversight': [
            'All 9.2.xxx Pre-Construction SOPs - Review and approval oversight',
            'All 9.41.xxx Execution SOPs - Performance monitoring',
            'All 9.6.xxx Closeout SOPs - Quality assurance review'
        ],
        'Management Procedures': [
            'Resource allocation procedures',
            'Project selection and go/no-go decisions',
            'Performance management procedures',
            'Escalation and resolution procedures',
            'Financial review procedures'
        ],
        'Business Development': [
            'Client relationship management',
            'Proposal review and approval',
            'Strategic planning participation'
        ]
    },
    'training_requirements': {
        'Executive Leadership (Required)': [
            'Strategic Thinking',
            'Presentation Skills',
            'Industry Leadership',
            'Business Development',
            'Financial Management for Leaders'
        ],
        'Operational Management (Required)': [
            'All Project Manager training requirements',
            'Advanced financial management',
            'Contract management',
            'Risk management'
        ],
        'People Development (Required)': [
            'Advanced performance management',
            'Succession planning',
            'Organizational development',
            'Change management'
        ]
    },
    'job_descriptions': [
        '7.3.2 Operations Manager - Primary role governed by this directive',
        'Branch Manager - Senior role governed by this directive',
        '7.3.1 Project Manager - Reports to Operations Manager',
        '7.2.5 Project Superintendent - Reports to Operations Manager'
    ],
    'compliance': [
        'Complete all required training within 12 months of promotion to branch management',
        'Conduct quarterly business reviews',
        'Submit monthly operational reports',
        'Complete annual strategic planning activities',
        'Maintain industry certifications and continuing education',
        'Participate in company leadership meetings and initiatives'
    ]
}

# Create remaining management directives
remaining_mds = [
    (md_9_4, 'Policy Manual 9.4 Management Directives for General Superintendents (Revised).docx'),
    (md_9_6, 'Policy Manual 9.6 Management Directives for Branch Management (Revised).docx'),
]

for md_data, filename in remaining_mds:
    create_revised_md(md_data, filename)

print()
print('=' * 60)
print('Remaining Revised Management Directives created with:')
print('  - Cross-Reference Summary Table')
print('  - Implementing SOPs section')
print('  - Training Requirements section')
print('  - Job Description Alignment section')
