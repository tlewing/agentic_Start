from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

output_dir = 'C:/Users/tewing/Desktop/Claude Projects/Revised Job Descriptions'
os.makedirs(output_dir, exist_ok=True)

def create_revised_job_description(jd_data, filename):
    """Create a revised Job Description with embedded cross-references"""
    doc = Document()

    # Title
    title = doc.add_heading(f"Policy Manual {jd_data['policy_number']}", level=1)
    doc.add_heading(jd_data['title'], level=2)

    # === CROSS-REFERENCE SUMMARY ===
    doc.add_heading('CROSS-REFERENCE SUMMARY', level=2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    row = table.rows[0]
    row.cells[0].text = 'Related SOPs'
    row.cells[1].text = jd_data.get('related_sops', 'N/A')

    row = table.rows[1]
    row.cells[0].text = 'Management Directives'
    row.cells[1].text = jd_data.get('management_directives', 'N/A')

    row = table.rows[2]
    row.cells[0].text = 'Required Training'
    row.cells[1].text = jd_data.get('required_training', 'N/A')

    row = table.rows[3]
    row.cells[0].text = 'Target Skill Rules'
    row.cells[1].text = jd_data.get('target_skill_rules', 'N/A')

    doc.add_paragraph()

    # Department and Reporting
    doc.add_paragraph(f"Department: {jd_data['department']}")
    doc.add_paragraph(f"Reports To: {jd_data['reports_to']}")
    doc.add_paragraph(f"Supervises: {jd_data['supervises']}")

    # Qualifications
    doc.add_heading('QUALIFICATIONS', level=2)
    for qual in jd_data['qualifications']:
        doc.add_paragraph(qual, style='List Bullet')

    # Summary
    doc.add_heading('SUMMARY', level=2)
    doc.add_paragraph(jd_data['summary'])

    # Essential Duties
    doc.add_heading('ESSENTIAL DUTIES AND RESPONSIBILITIES', level=2)
    for duty in jd_data['duties']:
        doc.add_paragraph(duty, style='List Bullet')

    # === NEW: Applicable SOPs Section ===
    doc.add_heading('APPLICABLE STANDARD OPERATING PROCEDURES', level=2)
    doc.add_paragraph('This role is responsible for executing the following SOPs:')
    for sop in jd_data.get('applicable_sops', []):
        doc.add_paragraph(sop, style='List Bullet')

    # === NEW: Required Training Section ===
    doc.add_heading('REQUIRED TRAINING', level=2)
    doc.add_paragraph('Personnel in this role must complete the following training per GSL Academy Target Skill Rules:')
    for training in jd_data.get('training_requirements', []):
        doc.add_paragraph(training, style='List Bullet')

    # === NEW: Management Directive Compliance ===
    doc.add_heading('MANAGEMENT DIRECTIVE COMPLIANCE', level=2)
    doc.add_paragraph(jd_data.get('md_compliance', 'This role must comply with applicable GSL Management Directives.'))

    # Skills and Competencies
    doc.add_heading('REQUIRED SKILLS AND COMPETENCIES', level=2)
    for skill in jd_data.get('skills', []):
        doc.add_paragraph(skill, style='List Bullet')

    # Physical Requirements (if applicable)
    if 'physical_requirements' in jd_data:
        doc.add_heading('PHYSICAL REQUIREMENTS', level=2)
        for req in jd_data['physical_requirements']:
            doc.add_paragraph(req, style='List Bullet')

    doc.save(f'{output_dir}/{filename}')
    print(f'Created: {filename}')

# ============================================
# JOB DESCRIPTIONS
# ============================================

# FOREMAN (7.2.4)
jd_foreman = {
    'policy_number': '7.2.4',
    'title': 'FOREMAN',
    'department': 'Operations',
    'reports_to': 'Project Manager / General Foreman',
    'supervises': '7-10 Field Employees (Journeymen, Apprentices)',

    # Cross-references
    'related_sops': '9.2.050 FS Reviews Plans; 9.41.470-510 Safety SOPs; 9.41.675-725 Field Operations SOPs; 9.41.215 Daily Reports; 9.3.xxx Site Setup SOPs',
    'management_directives': 'MD 9.3 (Field Leadership)',
    'required_training': 'Safety Coordinator Certification (1-11); Introduction to Field Leadership (all sections); Foreman Role & Responsibilities; Job Plan Setup; Performance Management; Lean Construction; Emotional Intelligence',
    'target_skill_rules': 'Foreman Safety Certification; Foreman Leadership Essentials; Foreman Operations Management; Foreman Performance Management; Foreman Lean Construction; Foreman Emotional Intelligence',

    'qualifications': [
        'Journeyman or Masters License or equivalent education and experience',
        'Certification as a GSL Safety Coordinator, including current certification in Standard First Aid and Adult CPR',
        'Current state issued Driver\'s License',
        'OSHA 30 (may be required in some states)',
        'Completion of GSL Foreman Training Program (Tabs 2-14, 22, 38, 41, 49, 57)',
        'Skill Level: L7 per GSL Academy Target Skill Rules'
    ],

    'summary': 'Directly supervise 7 to 10 employees in the Field Operations for the installation, maintenance and repair of electrical installations. Carry out supervisory responsibilities in accordance with the organization\'s policies and applicable laws. Responsibilities include training employees, planning, assigning, and directing work; appraising performance; rewarding and disciplining employees; addressing complaints and resolving problems. Primary responsibility is to bring the project in "Safely, On Time and within Budget".',

    'duties': [
        'All functions of a GSL Journeyman Electrician',
        'Completion of all requirements to achieve certification as a GSL Safety Coordinator including current certification in Standard First Aid and Adult CPR',
        'Production control, quality control and cost control over assigned work areas',
        'Implementation of lean construction principles',
        'Full implementation of GSL\'s Safety and Loss Control Program at the Foreman level',
        'Direct responsibility for crew\'s safety and enforcement of all company policies and procedures',
        'Training, protecting, directing, and managing Employee/Owners assigned to you',
        'Knowledge and utilization of all GSL standard operating procedures',
        'Plan, schedule, train, follow through and promote teamwork',
        'Maintain proper apprentice-to-journeyman ratios',
        'Weekly project updates to Project Manager',
        'Time card management and field records maintenance',
        'Change order identification and documentation',
        'Client relations development'
    ],

    'applicable_sops': [
        'SOP 9.2.050 - FS Reviews Plans, Specifications & Schedule',
        'SOP 9.2.055 - Compare Estimated vs. Planned Performance',
        'SOP 9.3.010-120 - Site Setup SOPs (Office, Storage, Power, etc.)',
        'SOP 9.41.470 - Develop Site-Specific Safety Plan',
        'SOP 9.41.475 - Conduct Safety Orientation',
        'SOP 9.41.480 - Conduct Safety Inspections',
        'SOP 9.41.485 - Conduct Safety Meetings',
        'SOP 9.41.490 - Report Safety Incidents',
        'SOP 9.41.675 - Conduct Daily Huddles',
        'SOP 9.41.680 - Conduct Weekly Foreman Meetings',
        'SOP 9.41.690 - Track Manpower Utilization',
        'SOP 9.41.695 - Track Production Quantities',
        'SOP 9.41.215 - Manage Daily Reports',
        'SOP 9.41.720 - Conduct Productivity Analysis',
        'SOP 9.6.005 - Prepare for Turnover (Field Perspective)'
    ],

    'training_requirements': [
        'Safety Coordinator Certification modules 1-11 (Tab 40) - Required before or upon promotion',
        'Introduction to Field Leadership Section 1 (Tab 51) - First 90 days',
        'Foreman Role & Responsibilities (Tab 51) - First 90 days',
        'Job Plan Setup and Maintenance (Tab 7) - First 6 months',
        'Performance Management Fundamentals (Tab 52) - First year',
        'Conducting Effective Performance Reviews (Tab 52) - First year',
        'Lean Construction Fundamentals (Tab 56) - First year',
        'Emotional Intelligence (Tab 51) - First year',
        'All safety training at Advanced level'
    ],

    'md_compliance': 'This role must comply with MD 9.3 (Management Directives for Field Leadership) which includes: maintaining positive attitude and self-discipline; pre-planning all projects with Project Manager; executing the plan by laying out daily work activities; monitoring productivity; completing timecards accurately; reviewing Job Cost vs Estimates monthly; never performing changed work without PM authorization; documenting schedule delays in daily reports.',

    'skills': [
        'Leadership and supervisory skills',
        'Production planning and scheduling',
        'Safety program implementation',
        'Quality control and inspection',
        'Cost control and budget management',
        'Employee training and development',
        'Effective communication (verbal and written)',
        'Problem resolution and decision-making',
        'Team building and motivation',
        'Time management',
        'Documentation and reporting'
    ],

    'physical_requirements': [
        'Ability to work in construction environment',
        'Standing, walking, climbing for extended periods',
        'Lifting up to 50 lbs regularly',
        'Working at heights and in confined spaces',
        'Exposure to weather conditions'
    ]
}

# PROJECT MANAGER (7.3.1)
jd_pm = {
    'policy_number': '7.3.1',
    'title': 'PROJECT MANAGER',
    'department': 'Operations',
    'reports_to': 'Branch Manager / Operations Manager',
    'supervises': 'Project Superintendent, Foremen, Project Coordinator',

    'related_sops': '9.2.xxx Pre-Construction SOPs; 9.41.xxx Execution SOPs; 9.6.xxx Closeout SOPs; All scheduling, documentation, change order SOPs',
    'management_directives': 'MD 9.5 (Project Management)',
    'required_training': 'Project Planning; Scheduling; Productivity Analysis; ProCore Fundamentals; ViewPoint Change Orders; Values-Driven Leadership; Strategic Thinking; Presentation Skills',
    'target_skill_rules': 'Project Manager Core; Project Manager Software; Project Manager Leadership',

    'qualifications': [
        'Bachelor\'s degree in Construction Management, Engineering, or related field (or equivalent experience)',
        '5+ years of electrical construction experience',
        '3+ years of project management experience',
        'Strong knowledge of NEC and electrical construction',
        'Proficiency in project management software (ProCore, ViewPoint, etc.)',
        'Current Driver\'s License'
    ],

    'summary': 'Manage all aspects of electrical construction projects from award through closeout. Responsible for project planning, scheduling, cost control, quality, safety, and client relations. Ensure projects are delivered safely, on time, and within budget while maintaining GSL quality standards and client satisfaction.',

    'duties': [
        'Lead project turnover meetings with Estimating',
        'Develop and manage project schedules',
        'Monitor and control project costs and budgets',
        'Coordinate with Field Leadership on resource allocation',
        'Manage change orders and contract modifications',
        'Ensure compliance with contract requirements',
        'Conduct regular project status meetings',
        'Prepare and submit progress billings',
        'Maintain project documentation in Plan Grid',
        'Interface with GC, Owner, and design team',
        'Support Field Leadership in problem resolution',
        'Ensure safety program compliance',
        'Conduct post-project reviews and capture lessons learned'
    ],

    'applicable_sops': [
        'SOP 9.2.005 - Review Contract for High-Risk Clauses',
        'SOP 9.2.010 - Team Selection',
        'SOP 9.2.015 - Project Turnover Meeting',
        'SOP 9.2.040 - PM Reviews Plans, Specifications & Schedule',
        'SOP 9.2.110 - Develop Project Schedule',
        'SOP 9.2.120 - Establish Tracking and Control Systems',
        'SOP 9.2.130 - Construction Execution Kickoff Meeting',
        'SOP 9.2.140 - Develop Project Budget',
        'SOP 9.41.315-350 - Scheduling SOPs',
        'SOP 9.41.355-390 - Scope and Change Order SOPs',
        'SOP 9.41.395-465 - Cost Control and Billing SOPs',
        'SOP 9.6.010-085 - Closeout SOPs'
    ],

    'training_requirements': [
        'Project Planning and Execution (Tab 2) - Upon promotion',
        'ProCore Fundamentals (Tab 57) - First 90 days',
        'ViewPoint Change Orders (Tab 8) - First 90 days',
        'Values-Driven Leadership (Tab 51) - Ongoing',
        'Scheduling Techniques (Tab 6) - First 6 months',
        'Productivity Analysis (Tab 56) - First 6 months',
        'Presentation Skills (Tab 51) - First year'
    ],

    'md_compliance': 'This role must comply with MD 9.5 (Management Directives for Project Management) which includes: contract scope review; schedule analysis; subcontract management; pre-planning requirements; prefabrication review; submittal and purchasing management; job monitoring with minimum twice monthly site visits; job plan updates; document control; change order management; weekly project status reports; progress billing procedures; and project closeout requirements.',

    'skills': [
        'Project management and planning',
        'Schedule development and management',
        'Cost control and financial analysis',
        'Contract administration',
        'Client relationship management',
        'Leadership and team management',
        'Problem-solving and decision-making',
        'Written and verbal communication',
        'Negotiation skills',
        'Software proficiency (ProCore, ViewPoint, Excel)'
    ]
}

# APPRENTICE ELECTRICIAN (7.2.1)
jd_apprentice = {
    'policy_number': '7.2.1',
    'title': 'APPRENTICE ELECTRICIAN',
    'department': 'Field Operations',
    'reports_to': 'Journeyman, Lead Journeyman, Foreman',
    'supervises': 'None',

    'related_sops': 'Safety SOPs (9.41.470-510); Tool Control SOPs; Site Work Requirements',
    'management_directives': 'MD 9.2 (Field Employees)',
    'required_training': 'New Hire Safety Orientation; NFPA 70E Fundamentals; Arc Flash Awareness; LOTO Fundamentals; Fall Protection; Ladder Safety; HazCom; Confined Space',
    'target_skill_rules': 'Apprentice Year 1 Onboarding; Apprentice Safety Fundamentals; Apprentice Field Operations',

    'qualifications': [
        'High school diploma or GED',
        'Enrollment in approved electrical apprenticeship program',
        'Valid Driver\'s License',
        'Ability to pass drug screening and background check',
        'Physical ability to perform electrical work'
    ],

    'summary': 'Complete 8,000 hours of on-the-job field experience and 144 hours of apprenticeship course work annually for 4 years. Perform field installation work under supervision of licensed Journeymen while achieving progressive skill levels across the 4-year program.',

    'duties': [
        'Complete all apprenticeship program requirements',
        'Follow safety policies of GSL, General Contractor, and Owner',
        'Perform electrical installation work under supervision',
        'Maintain tools and equipment properly',
        'Complete timecards accurately daily',
        'Report safety concerns immediately',
        'Accept mentoring and training from Journeymen',
        'Work as team player'
    ],

    'applicable_sops': [
        'SOP 9.41.475 - Safety Orientation (complete upon hire)',
        'SOP 9.41.485 - Safety Meetings (attend)',
        'All Safety SOPs (follow requirements)',
        'Site-specific work requirements'
    ],

    'training_requirements': [
        'New Hire Safety Orientation - First 7 days',
        'A Commitment to Zero Broken Lives - First 7 days',
        'NFPA 70E Fundamentals - First 90 days',
        'Arc Flash Awareness - First 90 days',
        'LOTO Fundamentals - First 90 days',
        'Fall Hazard Recognition - First 90 days',
        'Ladder Safety - First 90 days',
        'HazCom Program Fundamentals - First 90 days',
        '144 hours apprenticeship coursework annually'
    ],

    'md_compliance': 'This role must comply with MD 9.2 (Management Directives for Field Employees) which includes: following site work requirements; being signed in and ready at start time; accurate timecard completion; following break policies; staying with assigned work crew; cell phone policy compliance; PPE and dress code requirements; and signing acknowledgement of directives.',

    'skills': [
        'Basic electrical knowledge (developing)',
        'Blueprint reading (developing)',
        'Hand and power tool operation',
        'Safety awareness',
        'Communication skills',
        'Team collaboration',
        'Willingness to learn'
    ],

    'physical_requirements': [
        'Standing, walking, climbing for extended periods',
        'Lifting up to 50 lbs regularly',
        'Working at heights on ladders and scaffolds',
        'Working in confined spaces',
        'Exposure to weather conditions'
    ]
}

# JOURNEYMAN ELECTRICIAN (7.2.2)
jd_journeyman = {
    'policy_number': '7.2.2',
    'title': 'JOURNEYMAN ELECTRICIAN',
    'department': 'Operations',
    'reports_to': 'Foreman / General Foreman',
    'supervises': 'Apprentices (mentoring/training capacity)',

    'related_sops': 'All Apprentice SOPs plus Quality Control; Trade Coordination; Simple LOTO',
    'management_directives': 'MD 9.2 (Field Employees)',
    'required_training': 'All Apprentice Training plus Leadership Introduction; Mentoring & Coaching; Job Plan Fundamentals; Simple LOTO Procedures; Scaffold Fundamentals',
    'target_skill_rules': 'Journeyman Core Competencies; Journeyman Leadership Introduction; Journeyman Technical Advancement',

    'qualifications': [
        'Journeyman or Master Electrician License',
        'Completion of 4-year apprenticeship or equivalent',
        'Current state-issued Driver\'s License',
        'OSHA 10 (may be required in some states)',
        '8,000+ hours field experience'
    ],

    'summary': 'Perform installation, maintenance, and repair of electrical systems. Mentor and train apprentices. Execute field installation on projects as directed by Foreman. Obtain advanced technical knowledge in specialized areas. Primary focus on industrial and commercial electrical work.',

    'duties': [
        'All functions required of apprentice level',
        'Installation, maintenance, and repair of electrical systems',
        'Teaching, training, and mentoring apprentices',
        'Quality workmanship to GSL standards',
        'Proficiency in GSL standard operating procedures',
        'Implementation of lean construction principles',
        'Advanced technical work in specialized areas',
        'Prepare for management opportunities'
    ],

    'applicable_sops': [
        'All Apprentice-level SOPs',
        'SOP 9.41.525 - Quality Inspections (self-inspection)',
        'Simple LOTO Procedures',
        'Trade Coordination procedures'
    ],

    'training_requirements': [
        'All Apprentice training requirements (completed)',
        'Simple LOTO Procedures - Upon promotion',
        'Personal Fall Arrest Systems - Upon promotion',
        'Scaffold Fundamentals - Upon promotion',
        'Confined Space Program - Upon promotion',
        'Building Relationships - First year',
        'Mentoring & Coaching - First year',
        'Job Plan Fundamentals - First year'
    ],

    'md_compliance': 'This role must comply with MD 9.2 (Management Directives for Field Employees) with additional responsibility for quality workmanship and mentoring apprentices per GSL standards.',

    'skills': [
        'Advanced electrical installation skills',
        'Blueprint reading and interpretation',
        'NEC code knowledge',
        'Troubleshooting and problem-solving',
        'Mentoring and training ability',
        'Quality self-inspection',
        'Communication skills',
        'Team collaboration'
    ],

    'physical_requirements': [
        'Standing, walking, climbing for extended periods',
        'Lifting up to 50 lbs regularly',
        'Working at heights on ladders and scaffolds',
        'Working in confined spaces',
        'Exposure to weather conditions'
    ]
}

# PROJECT SAFETY COORDINATOR (7.2.6)
jd_safety_coord = {
    'policy_number': '7.2.6',
    'title': 'PROJECT SAFETY COORDINATOR',
    'department': 'Field Operations / Safety',
    'reports_to': 'Project Manager / Safety Director',
    'supervises': 'None (advisory/oversight role)',

    'related_sops': '9.41.470-510 All Safety SOPs; 9.2.070 Site Visit',
    'management_directives': 'MD 9.2; MD 9.3; MD 9.4 (Safety sections)',
    'required_training': 'Safety Coordinator modules 1-11; OSHA Focus Four; All Safety Training at Expert Level; First Aid; CPR',
    'target_skill_rules': 'Safety Coordinator Core; Safety Coordinator Training',

    'qualifications': [
        'Certification as GSL Safety Coordinator',
        'Current Standard First Aid Certification',
        'Adult CPR Certification',
        'OSHA 30 (strongly recommended/required in many states)',
        'Minimum 5 years electrical construction experience',
        'Strong knowledge of OSHA regulations and NFPA 70E'
    ],

    'summary': 'Implement and enforce GSL Safety and Loss Control Program on assigned projects. Conduct daily safety monitoring, incident prevention, safety meetings, OSHA compliance verification, incident investigation, and safety training coordination.',

    'duties': [
        'Implement GSL Safety and Loss Control Program',
        'Conduct daily safety monitoring and inspections',
        'Lead safety meetings and toolbox talks',
        'Ensure OSHA compliance',
        'Investigate and report incidents',
        'Coordinate safety training',
        'Conduct equipment and site safety inspections',
        'Coach employees on safety practices',
        'Maintain safety documentation'
    ],

    'applicable_sops': [
        'SOP 9.41.470 - Develop Site-Specific Safety Plan',
        'SOP 9.41.475 - Conduct Safety Orientation',
        'SOP 9.41.480 - Conduct Safety Inspections',
        'SOP 9.41.485 - Conduct Safety Meetings',
        'SOP 9.41.490 - Report Safety Incidents',
        'SOP 9.41.495 - Investigate Accidents',
        'SOP 9.41.500 - Conduct Emergency Preparedness',
        'SOP 9.41.505 - Conduct Safety Compliance Audits',
        'SOP 9.41.510 - Provide Safety Training'
    ],

    'training_requirements': [
        'Safety Coordinator modules 1-11 (Tab 40) - Required',
        'OSHA Focus Four Workplace Hazards (Tab 40) - Required',
        'All LOTO training at Expert level',
        'All Fall Protection training at Expert level',
        'Arc Flash and NFPA 70E at Expert level',
        'Incident Investigation training',
        'First Aid and CPR certification'
    ],

    'md_compliance': 'This role enforces MD 9.2 (safety requirements for field employees), MD 9.3 (safety management for field leadership), and MD 9.4 (safety enforcement for general superintendents). Primary responsibility for ensuring all workers comply with GSL Safety and Loss Control Program.',

    'skills': [
        'OSHA regulations expertise',
        'Safety program implementation',
        'Incident investigation',
        'Hazard assessment and mitigation',
        'Training and facilitation',
        'Communication and coaching',
        'Documentation and reporting',
        'NEC compliance knowledge'
    ]
}

# ESTIMATOR (7.3.3)
jd_estimator = {
    'policy_number': '7.3.3',
    'title': 'ESTIMATOR',
    'department': 'Operations',
    'reports_to': 'Department Manager / Branch Manager',
    'supervises': 'Level I Estimators (if Lead Estimator)',

    'related_sops': 'Project Turnover Meeting; Bid document management; Value engineering',
    'management_directives': 'MD 9.7 (Estimating)',
    'required_training': 'Accubid Fundamentals; Accubid Data Export; Job Plan Analysis; Excel Skills; Documentation Practices',
    'target_skill_rules': 'Estimator Level I Foundation; Estimator Standard Competencies',

    'qualifications': [
        'Field experience (preferred)',
        'Proficiency in Outlook, Word, and Excel',
        'Professional writing ability',
        'Blueprint reading capability',
        'NEC understanding',
        'OSHA 10 (developing to OSHA 30)'
    ],

    'summary': 'Prepare accurate and timely project estimates as directed by Branch Management. Estimates should reflect lowest cost means and methods that satisfy bid document requirements. Participate in project turnover meetings to transfer project knowledge to operations.',

    'duties': [
        'Prepare project bid estimates',
        'Perform material take-offs',
        'Calculate labor requirements',
        'Develop proposals with accurate pricing',
        'Submit timely RFIs during bid process',
        'Identify exclusions and clarifications',
        'Participate in project turnover meetings',
        'Review bid documents thoroughly',
        'Identify value engineering opportunities',
        'Coordinate with Branch Management on bid strategy'
    ],

    'applicable_sops': [
        'SOP 9.2.015 - Project Turnover Meeting (participate)',
        'SOP 9.2.057 - Identify VE and Prefabrication Opportunities',
        'Bid Document Management procedures',
        'RFI Process procedures'
    ],

    'training_requirements': [
        'Accubid Fundamentals (Tab 11) - First 90 days',
        'Excel Skills (Tab 34) - First 90 days',
        'Documentation Practices - First 90 days',
        'Accubid Data Export (Tab 11) - First year',
        'Job Plan Analysis - First year',
        'MD 9.7 Estimating Directives - Ongoing'
    ],

    'md_compliance': 'This role must comply with MD 9.7 (Management Directives for Estimating) which includes: systematic approach to each bid; finding unique advantages; using competitive labor units; knowing the competition; timely RFI submission; clear exclusions and clarifications; thorough document review; schedule analysis; value engineering alternates; and participation in project turnover meetings.',

    'skills': [
        'Electrical estimating',
        'Blueprint reading',
        'Material take-off',
        'Labor calculation',
        'Cost analysis',
        'Accubid software',
        'Excel proficiency',
        'Written communication',
        'Attention to detail'
    ]
}

# Create revised Job Descriptions
print('Creating Revised Job Descriptions with Cross-References...')
print('=' * 60)

create_revised_job_description(jd_foreman, 'Policy Manual 7.2.4 Foreman (Revised).docx')
create_revised_job_description(jd_pm, 'Policy Manual 7.3.1 Project Manager (Revised).docx')
create_revised_job_description(jd_apprentice, 'Policy Manual 7.2.1 Apprentice Electrician (Revised).docx')
create_revised_job_description(jd_journeyman, 'Policy Manual 7.2.2 Journeyman Electrician (Revised).docx')
create_revised_job_description(jd_safety_coord, 'Policy Manual 7.2.6 Project Safety Coordinator (Revised).docx')
create_revised_job_description(jd_estimator, 'Policy Manual 7.3.3 Estimator (Revised).docx')

print()
print('=' * 60)
print('Revised Job Descriptions created with:')
print('  - Cross-Reference Summary Table')
print('  - Applicable SOPs section')
print('  - Required Training section')
print('  - Management Directive Compliance section')
