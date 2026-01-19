from docx import Document
from docx.shared import Inches, Pt
import os
import sys

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

output_dir = 'C:/Users/tewing/Desktop/Claude Projects/Revised Management Directives'
os.makedirs(output_dir, exist_ok=True)

def create_revised_md(md_data, filename):
    """Create a revised Management Directive with embedded cross-references"""
    doc = Document()

    # Title
    doc.add_heading(f"Policy Manual {md_data['policy_number']}", level=1)
    doc.add_heading(md_data['title'], level=2)

    # === CROSS-REFERENCE SUMMARY ===
    doc.add_heading('CROSS-REFERENCE SUMMARY', level=2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    row = table.rows[0]
    row.cells[0].text = 'Target Audience'
    row.cells[1].text = md_data.get('target_audience', 'N/A')

    row = table.rows[1]
    row.cells[0].text = 'Related Job Descriptions'
    row.cells[1].text = md_data.get('job_descriptions', 'N/A')

    row = table.rows[2]
    row.cells[0].text = 'Implementing SOPs'
    row.cells[1].text = md_data.get('implementing_sops', 'N/A')

    row = table.rows[3]
    row.cells[0].text = 'Required Training'
    row.cells[1].text = md_data.get('required_training', 'N/A')

    doc.add_paragraph()

    # Purpose
    doc.add_heading('PURPOSE', level=2)
    doc.add_paragraph(md_data['purpose'])

    # Scope
    doc.add_heading('SCOPE', level=2)
    doc.add_paragraph(md_data['scope'])

    # Key Directives
    doc.add_heading('KEY DIRECTIVES', level=2)
    for section in md_data['directives']:
        doc.add_heading(section['name'], level=3)
        for directive in section['items']:
            doc.add_paragraph(directive, style='List Bullet')

    # === NEW: Implementing SOPs ===
    doc.add_heading('IMPLEMENTING STANDARD OPERATING PROCEDURES', level=2)
    doc.add_paragraph('The following SOPs implement the requirements of this Management Directive:')
    for sop in md_data.get('sop_list', []):
        doc.add_paragraph(sop, style='List Bullet')

    # === NEW: Training Requirements ===
    doc.add_heading('TRAINING REQUIREMENTS', level=2)
    doc.add_paragraph('Personnel subject to this directive must complete:')
    for training in md_data.get('training_list', []):
        doc.add_paragraph(training, style='List Bullet')

    # === NEW: Job Description Alignment ===
    doc.add_heading('JOB DESCRIPTION ALIGNMENT', level=2)
    doc.add_paragraph(md_data.get('jd_alignment', 'This directive applies to roles defined in the GSL Policy Manual.'))

    # Compliance
    doc.add_heading('COMPLIANCE AND ENFORCEMENT', level=2)
    doc.add_paragraph(md_data.get('compliance', 'Compliance with this directive is mandatory. Violations may result in disciplinary action.'))

    # Acknowledgement
    doc.add_heading('ACKNOWLEDGEMENT', level=2)
    doc.add_paragraph(md_data.get('acknowledgement', 'Personnel must sign acknowledgement of receipt and understanding of this directive.'))

    doc.save(f'{output_dir}/{filename}')
    print(f'Created: {filename}')

# ============================================
# MANAGEMENT DIRECTIVES
# ============================================

# MD 9.2 - Field Employees
md_9_2 = {
    'policy_number': '9.2',
    'title': 'Management Directives for Field Employees',

    'target_audience': 'GSL Field Employee-Owners (all field electricians and laborers)',
    'job_descriptions': 'Policy 7.2.1 (Apprentice); Policy 7.2.2 (Journeyman); Policy 7.2.3 (Lead Journeyman)',
    'implementing_sops': 'Safety SOPs (9.41.470-510); Site Work Requirements; Tool Control SOPs',
    'required_training': 'New Hire Safety Orientation; Safety Awareness & Culture; HazCom; Emergency Response; All role-specific safety training',

    'purpose': 'To establish clear expectations and requirements for all GSL Field Employee-Owners regarding work conduct, safety, attendance, and professional behavior on job sites.',

    'scope': 'Applies to all GSL field employees including apprentices, journeymen, and lead journeymen working on any GSL project.',

    'directives': [
        {'name': 'General Site Work Requirements',
         'items': [
             'Follow GSL\'s General Site Work Requirements and Site-Specific Work Requirements',
             'Follow all safety policies of GSL, General Contractor, and Owner',
             'Focus on safety and productive work',
             'Report safety concerns immediately to project leadership'
         ]},
        {'name': 'Work Schedule & Attendance',
         'items': [
             'Be signed in and ready at assigned start time with tools prepared',
             'Wear tool belts and bags during work shift',
             'Workday ends at designated work area; 10 minutes allowed for cleanup and timecards',
             'Accurately complete timecards daily and sign out',
             'Call designated branch call-in number and Foreman/Superintendent if late or absent',
             'Provide 2-week advance notice for time off (except emergencies)',
             'No call/no show results in disciplinary action'
         ]},
        {'name': 'Break & Rest Periods',
         'items': [
             'Two 10-minute breaks allowed during day',
             'Breaks taken at times minimizing workflow impact',
             '10-minute breaks must be taken on-site',
             'Lunch times scheduled based on project needs'
         ]},
        {'name': 'PPE & Dress Code',
         'items': [
             'Personal Protective Equipment (PPE) mandatory; violations not tolerated',
             'Follow GSL dress code and grooming standards',
             'No T-shirts or hard hats with logos of other electrical contractors',
             'Name label required on front of hard hat',
             'No smoking or vaping except in designated areas'
         ]},
        {'name': 'Cell Phone & Electronics',
         'items': [
             'No personal cell phones for surfing internet or gaming during work',
             'Cell phones never used while driving, operating equipment, or in safety hazard situations',
             'Personal calls only during lunch or after shift',
             'One earbud allowed if no safety violation (not hearing protection)'
         ]}
    ],

    'sop_list': [
        'SOP 9.41.475 - Conduct Safety Orientation (new hire requirement)',
        'SOP 9.41.485 - Conduct Safety Meetings (attendance required)',
        'SOP 9.41.490 - Report Safety Incidents (reporting requirement)',
        'All Safety SOPs - Follow requirements',
        'Tool Control SOPs - Tool responsibility',
        'Timecard procedures - Daily completion'
    ],

    'training_list': [
        'New Hire Safety Orientation - First 7 days',
        'A Commitment to Zero Broken Lives - First 7 days',
        'Safety Awareness & Culture - First 30 days',
        'HazCom Program Fundamentals - First 30 days',
        'Emergency Response - First 30 days',
        'Role-specific safety training per Target Skill Rules'
    ],

    'jd_alignment': 'This directive applies to all field positions defined in Policy Manual 7.2.1 (Apprentice Electrician), 7.2.2 (Journeyman Electrician), and 7.2.3 (Lead Journeyman). These job descriptions reference compliance with MD 9.2 requirements.',

    'compliance': 'Compliance with this directive is mandatory for all field employees. Violations may result in disciplinary action up to and including termination. General Superintendents are responsible for reviewing these directives with all field employees and collecting signed acknowledgements.',

    'acknowledgement': 'Newly hired employees must review and sign off during new hire orientation. All other field employees must also sign acknowledgement of these directives. Signed acknowledgements are maintained by Human Resources.'
}

# MD 9.3 - Field Leadership
md_9_3 = {
    'policy_number': '9.3',
    'title': 'Management Directives for Field Leadership',

    'target_audience': 'Project Foremen and Project Superintendents',
    'job_descriptions': 'Policy 7.2.4 (Foreman); Policy 7.2.5 (Project Superintendent)',
    'implementing_sops': '9.2.050 Plans Review; 9.41.xxx Field Operations SOPs; 9.41.215 Daily Reports; 9.41.675-680 Huddles and Meetings',
    'required_training': 'Introduction to Field Leadership (all sections); Foreman Role & Responsibilities; Job Plan Setup; Lean Construction; Performance Management',

    'purpose': 'To establish expectations and requirements for GSL Field Leadership regarding project execution, crew management, productivity, reporting, and compliance.',

    'scope': 'Applies to all Project Foremen and Project Superintendents responsible for leading field crews on GSL projects.',

    'directives': [
        {'name': 'Objectives',
         'items': [
             'Deliver projects on-time and within budget',
             'Achieve contractual quality requirements',
             'Meet regulatory safety requirements'
         ]},
        {'name': 'Attitude & Self-Discipline',
         'items': [
             'Demonstrate positive attitude about company, projects, clients, people, and work',
             'Show self-discipline to refrain from negative behaviors impacting morale',
             'Be ambassadors of good will for GSL',
             'Address concerns through proper channels (Project Superintendent/Project Manager)'
         ]},
        {'name': 'Pre-Planning',
         'items': [
             'All projects must be pre-planned with Project Manager',
             'Plan from start to finish',
             'Changes to plan handled more efficiently than on-the-fly planning'
         ]},
        {'name': 'Execution of the Plan',
         'items': [
             'Layout daily work activities for crews',
             'Monitor progress during work shift',
             'Report quantities for crews at end of shift',
             'If daily objectives not met, review issues and take corrective action',
             'Ensure labor, materials, tools, equipment available',
             'Quantities installed reported and entered in Job Plan'
         ]},
        {'name': 'Productivity',
         'items': [
             'Foremen and Lead Men accountable for crew productivity',
             'Direct all electricians to operate efficiently',
             'Wear tool belts and bags for anticipated work',
             'Proper material staging critical',
             'Report significant labor, material, equipment overage immediately to PM',
             'Focus on quality control and minimize mistakes/rework'
         ]},
        {'name': 'Reporting',
         'items': [
             'Complete employee timecards daily',
             'Record time worked to appropriate work codes',
             'Foreman timecards completed daily with hours and crew\'s installed quantities',
             'Accuracy in reporting time and quantity paramount'
         ]},
        {'name': 'Changed Work',
         'items': [
             'Never perform significant changed work without PM consultation and authorization',
             'Report concerns about change order delays to PM',
             'Never assume field changes are clarifications of base bid scope',
             'Only proceed with changed work when authorized by PM'
         ]},
        {'name': 'Daily Reports',
         'items': [
             'Submit Daily Reports with schedule slippage, impacts, productivity impacts',
             'Document work outside contract scope',
             'Information flows to PM for Weekly Project Status Reports',
             'Information is time-sensitive and critical for notice compliance'
         ]}
    ],

    'sop_list': [
        'SOP 9.2.050 - FS Reviews Plans, Specifications & Schedule',
        'SOP 9.41.675 - Conduct Daily Huddles',
        'SOP 9.41.680 - Conduct Weekly Foreman Meetings',
        'SOP 9.41.690 - Track Manpower Utilization',
        'SOP 9.41.695 - Track Production Quantities',
        'SOP 9.41.215 - Manage Daily Reports',
        'SOP 9.41.720 - Conduct Productivity Analysis',
        'SOP 9.41.470-510 - Safety Management SOPs',
        'All Site Setup SOPs (9.3.xxx)'
    ],

    'training_list': [
        'Introduction to Field Leadership Section 1 (Tab 51) - First 90 days',
        'Foreman Role & Responsibilities (Tab 51) - First 90 days',
        'Job Plan Setup and Maintenance (Tab 7) - First 6 months',
        'Leadership Essentials (Tab 51) - First 6 months',
        'Lean Construction Fundamentals (Tab 56) - First year',
        'Effective Communication (Tab 51) - First year',
        'Safety Coordinator Certification (Tab 40) - Required for Foreman'
    ],

    'jd_alignment': 'This directive applies to positions defined in Policy Manual 7.2.4 (Foreman) and 7.2.5 (Project Superintendent). These job descriptions include compliance with MD 9.3 requirements as an essential duty.',

    'compliance': 'Compliance is mandatory. Job Cost Review must occur at least monthly with Project Manager. Daily reports must be submitted timely. Failure to provide required notice and documentation can have severe financial impact on GSL.',

    'acknowledgement': 'Field Leadership must acknowledge receipt and understanding of these directives. General Superintendents are responsible for reviewing and obtaining acknowledgements.'
}

# MD 9.5 - Project Management
md_9_5 = {
    'policy_number': '9.5',
    'title': 'Management Directives for Project Management',

    'target_audience': 'Project Managers, Contract Managers, Business Managers',
    'job_descriptions': 'Policy 7.3.1 (Operations Management / Project Manager)',
    'implementing_sops': '9.2.xxx Pre-Construction SOPs; 9.41.xxx Execution SOPs; 9.6.xxx Closeout SOPs',
    'required_training': 'Project Planning; Scheduling; ProCore; ViewPoint; Values-Driven Leadership; Productivity Analysis',

    'purpose': 'To establish expectations and requirements for GSL Project Management regarding contract administration, pre-planning, execution monitoring, cost control, and project closeout.',

    'scope': 'Applies to all Project Managers, Contract Managers, and Business Managers responsible for GSL project delivery.',

    'directives': [
        {'name': 'Contract Review',
         'items': [
             'Scope of Work shall include work set forth in GSL\'s bid',
             'Scope must be consistent with Prime Agreement',
             'Review with Estimating to ensure clarifications identified in contract'
         ]},
        {'name': 'Schedule Analysis',
         'items': [
             'Perform schedule analysis prior to establishing bid price',
             'Request detailed copy of schedule if not included',
             'Determine cost impacts of aggressive or elongated schedules'
         ]},
        {'name': 'Pre-Planning',
         'items': [
             'Conduct Project Turnover Meetings with Estimating',
             'Review bid documents and pre-construction checklist',
             'Prepare Project Organization Chart',
             'Prepare man-loaded schedule and Schedule of Values'
         ]},
        {'name': 'Prefabrication',
         'items': [
             'Review each project for prefabrication opportunities',
             'Coordinate with GSL\'s prefabrication team',
             'Utilize BIM Modeling for raceway systems layout'
         ]},
        {'name': 'Scheduling',
         'items': [
             'Take initiative on project scheduling',
             'Work with Field Leadership on look ahead schedules',
             'Utilize Pull Planning by milestone',
             'Identify beginning dates and normal durations'
         ]},
        {'name': 'Job Monitoring',
         'items': [
             'Focus on executing the plan with sufficient job site time',
             'Minimum twice per month job site visits (every week if possible)',
             'Review daily reports and prepare Weekly Project Status Reports',
             'Communicate delays, disruptions, scope changes to Contract Managers'
         ]},
        {'name': 'Job Plans',
         'items': [
             'Track labor productivity by comparing budgeted to actual hours',
             'Determine positive/negative labor hour variances',
             'Update at least monthly for change orders and adjustments',
             'Share Job Costs/Estimates Report with Field Leadership weekly'
         ]},
        {'name': 'Change Orders',
         'items': [
             'Submit timely, accurate change order requests',
             'Follow-up on pending change orders',
             'Always request time extensions when necessary'
         ]},
        {'name': 'Project Closeout',
         'items': [
             'Complete Field Leadership Interviews',
             'Operations and Maintenance Manuals',
             'Letter of Warranty',
             'Final Progress and Retention Billing',
             'Employee Evaluations',
             'General Contractor or Owner Survey'
         ]}
    ],

    'sop_list': [
        'SOP 9.2.005 - Review Contract for High-Risk Clauses',
        'SOP 9.2.010 - Team Selection',
        'SOP 9.2.015 - Project Turnover Meeting',
        'SOP 9.2.040 - PM Reviews Plans, Specifications & Schedule',
        'SOP 9.2.110-130 - Schedule and Kickoff SOPs',
        'SOP 9.2.140 - Develop Project Budget',
        'SOP 9.41.315-350 - Scheduling SOPs',
        'SOP 9.41.355-390 - Change Order SOPs',
        'SOP 9.41.395-465 - Cost Control SOPs',
        'SOP 9.6.010-085 - Closeout SOPs'
    ],

    'training_list': [
        'Project Planning and Execution (Tab 2) - Upon promotion',
        'ProCore Fundamentals (Tab 57) - First 90 days',
        'ViewPoint Change Orders (Tab 8) - First 90 days',
        'Scheduling Techniques (Tab 6) - First 6 months',
        'Job Plan Management (Tab 7) - First 6 months',
        'Values-Driven Leadership (Tab 51) - Ongoing',
        'Productivity Analysis (Tab 56) - First year'
    ],

    'jd_alignment': 'This directive applies to Project Manager positions defined in Policy Manual 7.3.1. The Project Manager job description includes compliance with MD 9.5 as a primary responsibility.',

    'compliance': 'Compliance is mandatory. Weekly Project Status Reports must be submitted. Progress billings and job plan updates must be completed monthly by required dates. Under billings must be avoided.',

    'acknowledgement': 'Project Managers must acknowledge understanding of these directives upon assignment to project management role.'
}

# MD 9.7 - Estimating
md_9_7 = {
    'policy_number': '9.7',
    'title': 'Management Directives for Estimating',

    'target_audience': 'Estimating Department, Estimators, Lead Estimators',
    'job_descriptions': 'Policy 7.3.3 (Estimating)',
    'implementing_sops': 'Project Turnover Meeting; Bid document procedures; Value engineering',
    'required_training': 'Accubid Fundamentals; Excel Skills; NEC Training; MD 9.7 Directives',

    'purpose': 'To establish expectations and requirements for GSL Estimating Department regarding bid preparation, accuracy, strategy, and project turnover.',

    'scope': 'Applies to all Estimators and Lead Estimators responsible for preparing project bids.',

    'directives': [
        {'name': 'Primary Objective',
         'items': [
             'Produce accurate and timely estimates as directed by Branch Management',
             'Estimates should reflect lowest cost means and methods',
             'Satisfy bid document requirements'
         ]},
        {'name': 'Pre-Bid Strategy',
         'items': [
             'Systematically and strategically approach each bid',
             'Find at least three unique advantages for GSL on each project',
             'Work the bid, don\'t just go through the motions',
             'Search diligently for lowest cost means and methods',
             'Use competitive labor units',
             'Know the competition and determine their approach'
         ]},
        {'name': 'RFI Process',
         'items': [
             'Submit timely RFIs early in take-off process',
             'List any RFIs not responded to in bid proposal',
             'Identify actions taken regarding unanswered RFIs'
         ]},
        {'name': 'Exclusions & Clarifications',
         'items': [
             'All estimates must clearly state exclusions or clarifications',
             'If bid doesn\'t cover 100% of electrical work, clearly state exclusions',
             'Complete Estimating Checklist prior to bid review'
         ]},
        {'name': 'Document Review',
         'items': [
             'Account for all cost impacts in specifications',
             'Bound to all documents, all drawings, all specifications',
             'Thoroughly review other discipline documents for electrical scope',
             'If questions or uncomfortable bidding, ask for assistance'
         ]},
        {'name': 'Schedule Analysis',
         'items': [
             'Imperative to perform schedule analysis prior to bid price',
             'Request detailed schedule if not included',
             'Assess cost impacts for aggressive schedules'
         ]},
        {'name': 'Project Turnover',
         'items': [
             'Participate in Project Turnover Meeting',
             'Thoroughly review means and methods in take-off',
             'Review every exclusion/clarification in bid',
             'Review all prefabrication incorporated',
             'Discuss all major vendors and subcontractors'
         ]}
    ],

    'sop_list': [
        'SOP 9.2.015 - Project Turnover Meeting (participation required)',
        'SOP 9.2.057 - Identify VE and Prefabrication Opportunities',
        'Bid Document Management procedures',
        'RFI Process procedures',
        'Estimating Checklist completion'
    ],

    'training_list': [
        'Accubid Fundamentals (Tab 11) - First 90 days',
        'Accubid Data Export (Tab 11) - First year',
        'Excel Skills (Tab 34) - First 90 days',
        'NEC Training - Ongoing',
        'MD 9.7 Directives review - Upon hire'
    ],

    'jd_alignment': 'This directive applies to Estimator positions defined in Policy Manual 7.3.3. The Estimator job description includes compliance with MD 9.7 as a primary responsibility.',

    'compliance': 'Compliance is mandatory. Pre-bid evaluation forms required for bids over $2.0 million. Estimating Checklist must be completed prior to bid review. Never permissible to guess.',

    'acknowledgement': 'Estimators must acknowledge understanding of these directives upon joining the Estimating Department.'
}

# Create revised Management Directives
print('Creating Revised Management Directives with Cross-References...')
print('=' * 60)

create_revised_md(md_9_2, 'Policy Manual 9.2 Management Directives for Field Employees (Revised).docx')
create_revised_md(md_9_3, 'Policy Manual 9.3 Management Directives for Field Leadership (Revised).docx')
create_revised_md(md_9_5, 'Policy Manual 9.5 Management Directives for Project Management (Revised).docx')
create_revised_md(md_9_7, 'Policy Manual 9.7 Management Directives for Estimating (Revised).docx')

print()
print('=' * 60)
print('Revised Management Directives created with:')
print('  - Cross-Reference Summary Table')
print('  - Implementing SOPs section')
print('  - Training Requirements section')
print('  - Job Description Alignment section')
