from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import sys

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

output_dir = 'C:/Users/tewing/Desktop/Claude Projects/Revised SOPs'
os.makedirs(output_dir, exist_ok=True)

def create_revised_sop(sop_data, filename):
    """Create a revised SOP with embedded cross-references"""
    doc = Document()

    # Title
    title = doc.add_heading(f"SOP: {sop_data['number']} – {sop_data['title']}", level=1)

    # Department
    doc.add_paragraph(f"Department: {sop_data['department']}")

    # === NEW SECTION: Cross-Reference Summary ===
    doc.add_heading('CROSS-REFERENCE SUMMARY', level=2)

    # Create cross-reference table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    # Row 1: Related Job Descriptions
    row = table.rows[0]
    row.cells[0].text = 'Related Job Descriptions'
    row.cells[1].text = sop_data.get('job_descriptions', 'N/A')

    # Row 2: Management Directives
    row = table.rows[1]
    row.cells[0].text = 'Management Directives'
    row.cells[1].text = sop_data.get('management_directives', 'N/A')

    # Row 3: Related Training
    row = table.rows[2]
    row.cells[0].text = 'Related Training Modules'
    row.cells[1].text = sop_data.get('training_modules', 'N/A')

    # Row 4: Related SOPs
    row = table.rows[3]
    row.cells[0].text = 'Related SOPs'
    row.cells[1].text = sop_data.get('related_sops', 'N/A')

    doc.add_paragraph()

    # Purpose
    doc.add_heading('PURPOSE', level=2)
    doc.add_paragraph(sop_data['purpose'])

    # Scope
    doc.add_heading('SCOPE', level=2)
    doc.add_paragraph('Applies to all projects / situations where:')
    for scope_item in sop_data['scope']:
        doc.add_paragraph(scope_item, style='List Bullet')

    # Roles & Responsibilities
    doc.add_heading('ROLES & RESPONSIBILITIES', level=2)
    for role in sop_data['roles']:
        p = doc.add_paragraph()
        p.add_run(f"{role['title']} ({role['raci']})").bold = True
        for resp in role['responsibilities']:
            doc.add_paragraph(resp, style='List Bullet')

    doc.add_paragraph('RACI: R=Responsible, A=Accountable, C=Consulted, I=Informed')

    # Requirements
    doc.add_heading('REQUIREMENTS', level=2)
    for req in sop_data['requirements']:
        doc.add_paragraph(req, style='List Bullet')

    # Procedure
    doc.add_heading('PROCEDURE', level=2)
    for i, step in enumerate(sop_data['procedure'], 1):
        doc.add_heading(f"Step {i} – {step['name']}", level=3)
        doc.add_paragraph(step['description'])
        if 'substeps' in step:
            for substep in step['substeps']:
                doc.add_paragraph(substep, style='List Bullet')

    # Compliance & Review
    doc.add_heading('COMPLIANCE & REVIEW', level=2)
    doc.add_paragraph('KPIs:')
    for kpi in sop_data.get('kpis', []):
        doc.add_paragraph(kpi, style='List Bullet')

    # === NEW SECTION: Training Requirements ===
    doc.add_heading('TRAINING REQUIREMENTS', level=2)
    doc.add_paragraph('Personnel performing this SOP must complete the following training:')
    for training in sop_data.get('required_training', []):
        doc.add_paragraph(training, style='List Bullet')

    # === NEW SECTION: Management Directive Alignment ===
    doc.add_heading('MANAGEMENT DIRECTIVE ALIGNMENT', level=2)
    doc.add_paragraph(sop_data.get('md_alignment', 'This SOP aligns with GSL Management Directives.'))

    # Appendix
    doc.add_heading('APPENDIX', level=2)
    for appendix in sop_data.get('appendix', []):
        doc.add_paragraph(appendix, style='List Bullet')

    # Save
    doc.save(f'{output_dir}/{filename}')
    print(f'Created: {filename}')

# ============================================
# SOP 9.2.010 - Team Selection (Revised)
# ============================================
sop_9_2_010 = {
    'number': '9.2.010',
    'title': 'Team Selection',
    'department': 'Project Management / Operations',

    # NEW: Cross-references
    'job_descriptions': 'Policy 7.3.1 (Project Manager); Policy 7.2.4 (Foreman); Policy 7.2.5 (Project Superintendent); Policy 7.2.6 (Safety Coordinator)',
    'management_directives': 'MD 9.4 (Work Force Development - General Superintendents); MD 9.5 (Field Leadership & Manpower - Project Management)',
    'training_modules': 'Project Structure, Delegation, and Team Development; Chapter 4 Building an Effective Team; Introduction to Field Leadership Section 1',
    'related_sops': '9.2.015 - Project Turnover Meeting; 9.2.040 - PM Reviews Plans, Specs & Schedule; 9.2.130 - Construction Execution Kickoff Meeting',

    'purpose': 'To establish a standardized process for selecting and assembling the project team immediately after project award, ensuring the right mix of skills, experience, and availability to meet project requirements.',

    'scope': [
        'Project has been awarded',
        'Project team needs to be assembled',
        'Resource planning and assignment is required before kickoff'
    ],

    'roles': [
        {'title': 'Branch Manager', 'raci': 'A',
         'responsibilities': ['Selects Project Manager and assigns or approves final team assignments', 'Allocates resources across projects', 'Resolves resource conflicts']},
        {'title': 'Project Manager', 'raci': 'R',
         'responsibilities': ['Identifies required roles and skill sets', 'Recommends team members', 'Coordinates with Branch Manager on availability']},
        {'title': 'General Superintendent', 'raci': 'C',
         'responsibilities': ['Validates field leadership assignments', 'Confirms crew availability and capability']},
        {'title': 'Safety Coordinator', 'raci': 'C',
         'responsibilities': ['Ensures safety-qualified personnel are assigned', 'Reviews site-specific safety requirements']},
    ],

    'requirements': [
        'Award package and project scope',
        'Project schedule and duration',
        'Resource availability matrix',
        'Skills/certification requirements',
        'Project complexity assessment',
        'Historical performance data (similar projects)'
    ],

    'procedure': [
        {'name': 'Assess Project Requirements',
         'description': 'PM reviews award package and project scope to identify:',
         'substeps': ['Required roles (PM, APM, Foreman, Field Supervisor, Safety, etc.)', 'Skill requirements (certifications, experience level)', 'Duration and timing of resource needs', 'Special requirements (BIM, prefab, commissioning)']},
        {'name': 'Review Resource Availability',
         'description': 'PM coordinates with Branch Manager and GS to:',
         'substeps': ['Review current project commitments', 'Identify available personnel', 'Consider geographical and logistical factors', 'Assess workload balance across projects']},
        {'name': 'Select Team Members',
         'description': 'Based on requirements and availability:',
         'substeps': ['Match skills to project needs', 'Consider team chemistry and past performance', 'Ensure appropriate supervision ratios', 'Address any certification gaps']},
        {'name': 'Obtain Approval',
         'description': 'Branch Manager reviews and approves:',
         'substeps': ['Final team roster', 'Resource allocation timing', 'Any cross-project resource sharing', 'Budget implications of team selection']},
        {'name': 'Communicate Assignments',
         'description': 'PM communicates assignments to:',
         'substeps': ['Selected team members', 'Supporting departments (Safety, Purchasing, Prefab)', 'Project Coordinator for administrative setup', 'Schedule turnover meeting (SOP 9.2.015)']}
    ],

    'kpis': [
        'Team selection complete within 3 days of award - 100%',
        'All required roles identified and filled - 100%',
        'Branch Manager approval obtained - 100%',
        'Team members notified before turnover meeting - 100%'
    ],

    # NEW: Training Requirements
    'required_training': [
        'Project Manager: Project Structure, Delegation, and Team Development (Tab 51)',
        'Project Manager: Chapter 4 Building an Effective Team (Tab 51)',
        'General Superintendent: Introduction to Field Leadership Section 1 (Tab 51)',
        'All: Team Management training per Target Skill Rules'
    ],

    # NEW: MD Alignment
    'md_alignment': 'This SOP implements the requirements of MD 9.4 (Work Force Development Coordination) which directs General Superintendents to coordinate with Project Management and Field Leadership to assign desired field leadership and employees to projects. It also aligns with MD 9.5 (Field Leadership & Manpower) which requires Project Managers to coordinate with Branch Manager, Branch General Superintendent, and Branch Safety Superintendent to ensure best available field leadership and labor are assigned.',

    'appendix': [
        'Team Selection Checklist',
        'Resource Availability Matrix Template',
        'Skills/Certification Requirements by Role',
        'Project Complexity Assessment Form'
    ]
}

# ============================================
# SOP 9.2.015 - Project Turnover Meeting (Revised)
# ============================================
sop_9_2_015 = {
    'number': '9.2.015',
    'title': 'Project Turnover Meeting',
    'department': 'Project Management / Estimating / Operations',

    'job_descriptions': 'Policy 7.3.1 (Project Manager); Policy 7.3.3 (Estimator); Policy 7.2.4 (Foreman); Policy 7.2.5 (Project Superintendent)',
    'management_directives': 'MD 9.5 (Project Pre-Planning & Construction Plan Development); MD 9.7 (Project Turnover Meeting)',
    'training_modules': 'Introduction to Field Leadership Section 1; Chapter 0 Introduction to Critical Leadership Training; Introduction to Project Planning and Execution',
    'related_sops': '9.2.010 - Team Selection; 9.2.040 - PM Reviews Plans, Specs & Schedule; 9.2.050 - FS Reviews Plans, Specs & Schedule; 9.2.130 - Construction Execution Kickoff Meeting',

    'purpose': 'To establish a standardized process for conducting the project turnover meeting between Estimating and Project Management/Field Leadership, ensuring complete transfer of project knowledge, means and methods, and bid assumptions.',

    'scope': [
        'Project has been awarded and team selected',
        'Turnover from Estimating to Operations is required',
        'Project planning phase is beginning'
    ],

    'roles': [
        {'title': 'Project Manager', 'raci': 'R',
         'responsibilities': ['Schedules and facilitates turnover meeting', 'Documents turnover items and action items', 'Ensures all questions are addressed']},
        {'title': 'Estimator/Lead Estimator', 'raci': 'R',
         'responsibilities': ['Presents bid scope, means and methods', 'Reviews all exclusions and clarifications', 'Explains prefabrication assumptions', 'Discusses major vendors and subcontractors']},
        {'title': 'Field Supervisor', 'raci': 'C',
         'responsibilities': ['Reviews field installation approach', 'Identifies potential field issues', 'Confirms understanding of scope']},
        {'title': 'Branch Manager', 'raci': 'A',
         'responsibilities': ['Ensures turnover meeting occurs', 'Resolves any scope or resource issues']},
    ],

    'requirements': [
        'Award package and contract documents',
        'Complete bid documents in estimating folder',
        'Bid proposal with exclusions and clarifications',
        'Project Turnover Meeting Agenda',
        'Project Turnover Checklist',
        'Contract Summary'
    ],

    'procedure': [
        {'name': 'Schedule Meeting',
         'description': 'PM schedules turnover meeting within 5 days of team selection:',
         'substeps': ['Invite PM, Estimator, Field Supervisor, Branch Manager', 'Reserve adequate time (typically 2-4 hours)', 'Prepare meeting agenda']},
        {'name': 'Review Bid Documents',
         'description': 'Estimator presents comprehensive review:',
         'substeps': ['All bid documents identified and numbered', 'Scope of work and contract documents', 'All exclusions and clarifications in bid', 'RFIs submitted and responses received']},
        {'name': 'Review Means and Methods',
         'description': 'Estimator explains installation approach:',
         'substeps': ['Means and methods utilized in take-off', 'Prefabrication incorporated', 'Value engineering options', 'Labor units and assumptions']},
        {'name': 'Review Major Contracts',
         'description': 'Discuss vendors and subcontractors:',
         'substeps': ['Major vendor commitments', 'Subcontractor scopes', 'Pricing commitments and expirations', 'Material lead times']},
        {'name': 'Review Schedule and Resources',
         'description': 'Confirm schedule and manpower:',
         'substeps': ['Construction schedule and interim milestones', 'GSL field leadership and manpower requirements', 'Resource loading assumptions', 'Critical path items']},
        {'name': 'Document and Distribute',
         'description': 'PM documents all turnover items:',
         'substeps': ['Complete Project Turnover Checklist', 'Document action items with owners', 'Distribute Contract Summary', 'Move bid documents to job folder and Plan Grid']}
    ],

    'kpis': [
        'Turnover meeting held within 5 days of team selection - 100%',
        'Project Turnover Checklist completed - 100%',
        'All exclusions/clarifications documented - 100%',
        'Bid documents transferred to Plan Grid - 100%'
    ],

    'required_training': [
        'Project Manager: Introduction to Project Planning and Execution (Tab 2)',
        'Estimator: MD 9.7 Estimating Directives training',
        'Field Supervisor: Introduction to Field Leadership Section 1 (Tab 51)',
        'All: Chapter 0 Introduction to Critical Leadership Training'
    ],

    'md_alignment': 'This SOP implements MD 9.5 (Project Pre-Planning) which requires Project Turnover Meetings where Contract Manager, Branch Manager, Project Manager, and Field Leadership meet with Estimators. It also implements MD 9.7 (Project Turnover Meeting) which requires Estimating Department participation following the Turn-Over Meeting Agenda, including thorough review of means and methods, exclusions, prefabrication, value engineering, vendors, subcontractors, schedule, and Contract Summary.',

    'appendix': [
        'Project Turnover Meeting Agenda',
        'Project Turnover Checklist',
        'Contract Summary Template',
        'Bid Document Log Template'
    ]
}

# ============================================
# SOP 9.41.470 - Develop Site-Specific Safety Plan (Revised)
# ============================================
sop_9_41_470 = {
    'number': '9.41.470',
    'title': 'Develop Site-Specific Safety Plan',
    'department': 'Safety / Field Operations',

    'job_descriptions': 'Policy 7.2.6 (Project Safety Coordinator); Policy 7.4.6 (Safety and Loss Control Manager); Policy 7.2.4 (Foreman)',
    'management_directives': 'MD 9.2 (Safety Requirements - Field Employees); MD 9.3 (Safety Management - Field Leadership); MD 9.4 (Safety Enforcement - General Superintendents)',
    'training_modules': 'Safety Coordinator 1-11 (all modules); Introduction to Safety Management; OSHA Focus Four Workplace Hazards; New Hire Safety Orientation; A Commitment to Zero Broken Lives',
    'related_sops': '9.41.475 - Conduct Safety Orientation; 9.41.480 - Conduct Safety Inspections; 9.41.485 - Conduct Safety Meetings; 9.2.070 - Conduct Site Visit',

    'purpose': 'To establish a standardized process for developing a comprehensive site-specific safety plan that addresses all hazards unique to the project site and ensures compliance with GSL Safety and Loss Control Program, OSHA requirements, and client/GC safety requirements.',

    'scope': [
        'New project mobilization',
        'Significant change in project scope or conditions',
        'Client or regulatory requirement for site-specific plan'
    ],

    'roles': [
        {'title': 'Project Safety Coordinator', 'raci': 'R',
         'responsibilities': ['Develops site-specific safety plan', 'Conducts hazard assessment', 'Coordinates with client/GC safety requirements']},
        {'title': 'Safety and Loss Control Manager', 'raci': 'A',
         'responsibilities': ['Reviews and approves safety plan', 'Ensures regulatory compliance', 'Provides resources and support']},
        {'title': 'Project Manager', 'raci': 'C',
         'responsibilities': ['Provides project scope and schedule', 'Ensures safety plan is incorporated into project plan', 'Allocates safety resources']},
        {'title': 'Foreman', 'raci': 'C',
         'responsibilities': ['Provides field input on hazards', 'Implements safety plan requirements', 'Conducts daily safety activities']},
    ],

    'requirements': [
        'Project scope and work activities',
        'Site conditions and hazard assessment',
        'Client/GC safety requirements',
        'OSHA regulations applicable to work',
        'GSL Safety and Loss Control Program manual',
        'Equipment and PPE requirements'
    ],

    'procedure': [
        {'name': 'Conduct Hazard Assessment',
         'description': 'Safety Coordinator identifies all site hazards:',
         'substeps': ['Review project scope and work activities', 'Conduct site visit per SOP 9.2.070', 'Identify OSHA Focus Four hazards (falls, struck-by, caught-in, electrical)', 'Document all identified hazards']},
        {'name': 'Review Regulatory Requirements',
         'description': 'Identify all applicable regulations:',
         'substeps': ['OSHA construction standards', 'NFPA 70E electrical safety', 'State and local requirements', 'Client/GC specific requirements']},
        {'name': 'Develop Safety Plan',
         'description': 'Create comprehensive safety plan including:',
         'substeps': ['Emergency action plan', 'Hazard-specific controls', 'PPE requirements by task', 'Safety meeting schedule', 'Inspection and audit schedule', 'Incident reporting procedures']},
        {'name': 'Review and Approve',
         'description': 'Obtain required approvals:',
         'substeps': ['Safety Manager review and approval', 'Project Manager acknowledgment', 'Client/GC approval if required', 'Document approvals']},
        {'name': 'Communicate and Implement',
         'description': 'Roll out safety plan to team:',
         'substeps': ['Conduct safety orientation per SOP 9.41.475', 'Post safety plan in job trailer', 'Distribute to all supervision', 'Begin safety meeting schedule per SOP 9.41.485']}
    ],

    'kpis': [
        'Safety plan completed before mobilization - 100%',
        'All identified hazards addressed - 100%',
        'Safety Manager approval obtained - 100%',
        'All workers oriented to plan - 100%'
    ],

    'required_training': [
        'Safety Coordinator: Safety Coordinator modules 1-11 (Tab 40)',
        'Safety Coordinator: OSHA Focus Four Workplace Hazards (Tab 40)',
        'Foreman: Introduction to Safety Management (Tab 40)',
        'Foreman: GSL Safety Coordinator Certification',
        'All Field: New Hire Safety Orientation; A Commitment to Zero Broken Lives'
    ],

    'md_alignment': 'This SOP implements MD 9.2 (Safety Requirements) which requires all field employees to follow GSL Safety Policies and report safety concerns. It implements MD 9.3 (Safety Management) which requires Field Leadership to implement GSL Safety and Loss Control Program. It also aligns with MD 9.4 which requires General Superintendents to effectively enforce GSL Safety Policies and Procedures.',

    'appendix': [
        'Site-Specific Safety Plan Template',
        'Hazard Assessment Checklist',
        'Emergency Action Plan Template',
        'PPE Matrix by Task',
        'Safety Meeting Schedule Template'
    ]
}

# ============================================
# SOP 9.41.675 - Conduct Daily Huddles (Revised)
# ============================================
sop_9_41_675 = {
    'number': '9.41.675',
    'title': 'Conduct Daily Huddles',
    'department': 'Field Operations',

    'job_descriptions': 'Policy 7.2.4 (Foreman); Policy 7.2.3 (Lead Journeyman); Policy 7.2.5 (Project Superintendent)',
    'management_directives': 'MD 9.3 (Execution of the Plan - Field Leadership)',
    'training_modules': 'Techniques for Better Communication and Feedback; Chapter 8 Introduction to Effective Communication; Leadership Essentials; Building and Maintaining Team Morale',
    'related_sops': '9.41.680 - Conduct Weekly Foreman Meetings; 9.41.215 - Manage Daily Reports; 9.41.485 - Conduct Safety Meetings',

    'purpose': 'To establish a standardized process for conducting daily crew huddles to communicate work assignments, safety reminders, coordinate activities, and ensure all crew members understand daily objectives.',

    'scope': [
        'Start of each work shift',
        'After lunch breaks as needed',
        'When work conditions or plans change significantly'
    ],

    'roles': [
        {'title': 'Foreman', 'raci': 'R',
         'responsibilities': ['Conducts daily huddle', 'Communicates work assignments', 'Reviews safety topics', 'Addresses crew questions']},
        {'title': 'Lead Journeyman', 'raci': 'C',
         'responsibilities': ['Assists with huddle coordination', 'Provides input on crew assignments', 'Supports Foreman communication']},
        {'title': 'Project Superintendent', 'raci': 'A',
         'responsibilities': ['Ensures huddles occur', 'Provides project-level updates', 'Addresses escalated issues']},
    ],

    'requirements': [
        'Daily work plan and assignments',
        'Safety topic for the day',
        'Coordination needs with other trades',
        'Material and tool availability status',
        'Schedule milestones and priorities'
    ],

    'procedure': [
        {'name': 'Prepare for Huddle',
         'description': 'Foreman prepares before crew arrival:',
         'substeps': ['Review daily work plan from Job Plan', 'Identify safety topic (JHA items)', 'Note coordination requirements', 'Identify material/tool needs']},
        {'name': 'Conduct Huddle',
         'description': 'At shift start, gather crew and cover:',
         'substeps': ['Safety moment or JHA review', 'Daily work assignments by crew member', 'Production targets and quality reminders', 'Coordination with other trades', 'Material staging and tool locations', 'Questions and concerns']},
        {'name': 'Monitor and Adjust',
         'description': 'Throughout the day:',
         'substeps': ['Monitor progress against plan', 'Address issues as they arise', 'Conduct brief afternoon huddle if needed', 'Note productivity and quantities']},
        {'name': 'Document',
         'description': 'Record daily activities:',
         'substeps': ['Complete timecard with hours by cost code', 'Record installed quantities', 'Note any delays or impacts', 'Prepare for next day huddle']}
    ],

    'kpis': [
        'Daily huddle conducted - 100%',
        'Safety topic covered - 100%',
        'All crew members present at huddle - 95%',
        'Work assignments communicated clearly - 100%'
    ],

    'required_training': [
        'Foreman: Leadership Essentials (Tab 51)',
        'Foreman: Techniques for Better Communication and Feedback (Tab 51)',
        'Foreman: Chapter 8 Introduction to Effective Communication (Tab 51)',
        'Lead Journeyman: Building and Maintaining Team Morale (Tab 51)'
    ],

    'md_alignment': 'This SOP implements MD 9.3 (Execution of the Plan) which requires Field Leadership to layout daily work activities for crews, monitor progress during work shift, and report quantities for crews at end of shift. It ensures Foremen and Lead Men are accountable for crew productivity and direct all electricians to operate efficiently.',

    'appendix': [
        'Daily Huddle Checklist',
        'Safety Topic Calendar',
        'Work Assignment Template',
        'Coordination Log'
    ]
}

# ============================================
# SOP 9.6.070 - Conduct Post-Project Review (Revised)
# ============================================
sop_9_6_070 = {
    'number': '9.6.070',
    'title': 'Conduct Post-Project Review',
    'department': 'Project Management / Operations',

    'job_descriptions': 'Policy 7.3.1 (Project Manager); Policy 7.2.4 (Foreman); Policy 7.2.5 (Project Superintendent); Policy 7.3.3 (Estimator)',
    'management_directives': 'MD 9.5 (Project Closeout - Required Documents)',
    'training_modules': 'Continuous Improvement; Leadership Essentials; Chapter 15 Continuous Improvement',
    'related_sops': '9.6.085 - Document Lessons Learned; 9.6.080 - Capture Client Feedback; 9.6.010 - Conduct Closeout Meetings',

    'purpose': 'To establish a standardized process for conducting comprehensive post-project reviews to capture lessons learned, evaluate project performance, and drive continuous improvement across the organization.',

    'scope': [
        'Project substantial completion achieved',
        'Final billing submitted',
        'All closeout documentation complete'
    ],

    'roles': [
        {'title': 'Project Manager', 'raci': 'R',
         'responsibilities': ['Schedules and facilitates review', 'Prepares performance data', 'Documents lessons learned']},
        {'title': 'Branch Manager', 'raci': 'A',
         'responsibilities': ['Ensures reviews occur', 'Reviews performance results', 'Implements improvement actions']},
        {'title': 'Field Supervisor/Foreman', 'raci': 'C',
         'responsibilities': ['Provides field perspective', 'Identifies what worked well', 'Suggests improvements']},
        {'title': 'Estimator', 'raci': 'C',
         'responsibilities': ['Compares bid to actual', 'Provides estimating insights', 'Updates labor units database']},
    ],

    'requirements': [
        'Final job cost report',
        'Original estimate vs actual comparison',
        'Schedule performance data',
        'Safety record summary',
        'Quality metrics',
        'Client feedback (SOP 9.6.080)'
    ],

    'procedure': [
        {'name': 'Schedule Review Meeting',
         'description': 'PM schedules within 30 days of substantial completion:',
         'substeps': ['Invite key project team members', 'Include Estimating representative', 'Reserve adequate time (1-2 hours)', 'Gather performance data']},
        {'name': 'Review Project Performance',
         'description': 'Analyze key metrics:',
         'substeps': ['Budget vs actual (labor, material, equipment)', 'Schedule performance (milestones, completion)', 'Safety record (incidents, near misses)', 'Quality metrics (rework, punch list)', 'Client satisfaction']},
        {'name': 'Identify Lessons Learned',
         'description': 'Discuss what worked and what to improve:',
         'substeps': ['Pre-construction planning effectiveness', 'Execution challenges and solutions', 'Coordination and communication', 'Estimating accuracy', 'Field productivity']},
        {'name': 'Document and Distribute',
         'description': 'Capture and share lessons learned:',
         'substeps': ['Complete Lessons Learned form (SOP 9.6.085)', 'Update estimating database if applicable', 'Distribute to Branch Management', 'Share relevant lessons with organization']},
        {'name': 'Implement Improvements',
         'description': 'Take action on findings:',
         'substeps': ['Assign action items for process improvements', 'Update SOPs if needed', 'Incorporate into future project planning', 'Recognize team successes']}
    ],

    'kpis': [
        'Post-project review conducted within 30 days - 100%',
        'Lessons learned documented - 100%',
        'Action items assigned with owners - 100%',
        'Results shared with organization - 100%'
    ],

    'required_training': [
        'Project Manager: Continuous Improvement (Tab 56)',
        'Project Manager: Chapter 15 Continuous Improvement (Tab 56)',
        'All: Leadership Essentials (Tab 51)'
    ],

    'md_alignment': 'This SOP implements MD 9.5 (Project Closeout - Required Documents) which lists Field Leadership Interviews and Employee Evaluations as required closeout items. The post-project review process supports continuous improvement aligned with Lean Construction principles and helps improve future project performance.',

    'appendix': [
        'Post-Project Review Agenda',
        'Performance Metrics Summary Template',
        'Lessons Learned Form',
        'Estimating Feedback Form'
    ]
}

# Create the revised SOPs
print('Creating Revised SOPs with Cross-References...')
print('=' * 60)

create_revised_sop(sop_9_2_010, 'SOP 9.2.010 - Team Selection (Revised).docx')
create_revised_sop(sop_9_2_015, 'SOP 9.2.015 - Project Turnover Meeting (Revised).docx')
create_revised_sop(sop_9_41_470, 'SOP 9.41.470 - Develop Site-Specific Safety Plan (Revised).docx')
create_revised_sop(sop_9_41_675, 'SOP 9.41.675 - Conduct Daily Huddles (Revised).docx')
create_revised_sop(sop_9_6_070, 'SOP 9.6.070 - Conduct Post-Project Review (Revised).docx')

print()
print('=' * 60)
print('Revised SOPs created with embedded cross-references:')
print('  - Related Job Descriptions')
print('  - Management Directives')
print('  - Related Training Modules')
print('  - Training Requirements section')
print('  - Management Directive Alignment section')
